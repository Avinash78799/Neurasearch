import json
import re
import logging
from io import BytesIO
from typing import List, Dict, Any
from fpdf import FPDF
from docx import Document as DocxDocument

from database import db
from workspace_service import WorkspaceContext
from core.exceptions import KnowledgeError
from core.model_registry import get_llm

logger = logging.getLogger("neurasearch.knowledge_page")

# Prompts for AI note organization
ORGANIZE_PROMPT = """You are an expert information architect. Given the following knowledge items belonging to a single page, analyze their titles and summaries, and determine the most logical, structured order for them to flow.

Items list:
{items_json}

Your response must be a single valid JSON array of strings containing ONLY the item IDs in the suggested logical order (e.g. ["id_b", "id_a", "id_c"]). Do not include any introductory explanation or other content.
"""


class KnowledgePageService:
    """Manages knowledge page assembly, reordering, AI layout assistance, and multi-format exports."""

    @staticmethod
    def _verify_page_ownership(page_id: str, context: WorkspaceContext) -> Dict[str, Any]:
        """Ensures the page exists and belongs to the workspace context."""
        page_row = db.get_knowledge_item(page_id, context)
        if not page_row:
            raise KnowledgeError(f"Knowledge page '{page_id}' not found.", "PAGE_NOT_FOUND")
        page = dict(page_row)
        if page.get("type") != "page":
            raise KnowledgeError(f"Knowledge item '{page_id}' is not of type 'page'.", "INVALID_ITEM_TYPE")
        return page

    @staticmethod
    def get_references(page_id: str, context: WorkspaceContext) -> List[Dict[str, Any]]:
        """Fetch referenced items in order for a page."""
        KnowledgePageService._verify_page_ownership(page_id, context)
        return db.get_knowledge_page_items(page_id, context)

    @staticmethod
    def add_reference(page_id: str, item_id: str, position: int, context: WorkspaceContext) -> None:
        """Add a knowledge item reference to a page, verifying workspace isolation and loop safety."""
        # 1. Verify page ownership
        KnowledgePageService._verify_page_ownership(page_id, context)
        
        # 2. Prevent self-reference loops
        if page_id == item_id:
            raise KnowledgeError("A knowledge page cannot reference itself.", "SELF_REFERENCE_PROHIBITED")
            
        # 3. Verify target item exists in same workspace
        target = db.get_knowledge_item(item_id, context)
        if not target:
            raise KnowledgeError(f"Referenced knowledge item '{item_id}' not found.", "REFERENCED_ITEM_NOT_FOUND")

        # 4. Insert database reference row
        db.add_knowledge_page_item(page_id, item_id, position)
        logger.info("Added reference: page %s -> item %s (position %d)", page_id, item_id, position)

    @staticmethod
    def remove_reference(page_id: str, item_id: str, context: WorkspaceContext) -> bool:
        """Remove a reference relation from a page."""
        # Verify page ownership
        KnowledgePageService._verify_page_ownership(page_id, context)
        return db.remove_knowledge_page_item(page_id, item_id)

    @staticmethod
    def reorder_references(page_id: str, item_ids: List[str], context: WorkspaceContext) -> None:
        """Reorder all referenced items for a page."""
        # 1. Verify page ownership
        KnowledgePageService._verify_page_ownership(page_id, context)
        
        # 2. Verify all target items exist in active workspace
        for item_id in item_ids:
            if page_id == item_id:
                raise KnowledgeError("A knowledge page cannot reference itself.", "SELF_REFERENCE_PROHIBITED")
            target = db.get_knowledge_item(item_id, context)
            if not target:
                raise KnowledgeError(f"Referenced knowledge item '{item_id}' not found.", "REFERENCED_ITEM_NOT_FOUND")
                
        # 3. Run reorder query
        db.reorder_knowledge_page_items(page_id, item_ids)
        logger.info("Reordered references for page %s: %s", page_id, item_ids)

    @staticmethod
    async def ai_organize_page(page_id: str, context: WorkspaceContext) -> List[str]:
        """Suggests a logical ordered list of item IDs using Ollama (Temp=0.2)."""
        # 1. Fetch current items
        items = KnowledgePageService.get_references(page_id, context)
        if not items:
            return []
            
        # 2. Format basic metadata for prompt
        meta_items = []
        for item in items:
            meta_items.append({
                "id": item["id"],
                "title": item["title"],
                "summary": item.get("summary") or item["content"][:200]
            })
            
        llm = get_llm()
        prompt = ORGANIZE_PROMPT.format(items_json=json.dumps(meta_items, indent=2))
        
        try:
            response = await llm.ainvoke(prompt)
            raw = response.content.strip()
            
            # Extract JSON array
            match = re.search(r"\[.*\]", raw, re.DOTALL)
            if match:
                raw = match.group(0)
                
            suggested_ids = json.loads(raw)
            if not isinstance(suggested_ids, list):
                raise ValueError()
                
            # Filter suggested IDs to ensure they only contain valid references
            valid_ids = [item["id"] for item in items]
            ordered_ids = [str(x) for x in suggested_ids if x in valid_ids]
            
            # Append any missing IDs that LLM omitted
            for v_id in valid_ids:
                if v_id not in ordered_ids:
                    ordered_ids.append(v_id)
            return ordered_ids
        except Exception as exc:
            logger.warning("AI page organize suggestion failed: %s. Raw LLM content: %s", exc, raw if 'raw' in locals() else "")
            # Return current list of IDs on failure
            return [item["id"] for item in items]

    @staticmethod
    def export_markdown(page_id: str, context: WorkspaceContext) -> str:
        """Stitches the page and all referenced notes into a single Markdown document string."""
        page = KnowledgePageService._verify_page_ownership(page_id, context)
        items = KnowledgePageService.get_references(page_id, context)
        
        lines = []
        lines.append(f"# {page['title']}")
        if page.get("summary"):
            lines.append(f"\n> {page['summary']}\n")
        else:
            lines.append("")
            
        if page["content"]:
            lines.append(page["content"])
            lines.append("")
            
        lines.append("---")
        lines.append("## Referenced Notes & Insights")
        lines.append("")
        
        for idx, item in enumerate(items, 1):
            lines.append(f"### {idx}. {item['title']}")
            lines.append(f"*Type: {item['type'].capitalize()}*")
            lines.append("")
            lines.append(item["content"])
            lines.append("")
            lines.append("---")
            
        return "\n".join(lines)

    @staticmethod
    def export_pdf(page_id: str, context: WorkspaceContext) -> bytes:
        """Generates a clean PDF document from the assembled page elements using FPDF2."""
        page = KnowledgePageService._verify_page_ownership(page_id, context)
        items = KnowledgePageService.get_references(page_id, context)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Clean helper for built-in font safety
        def safe_txt(s: str) -> str:
            if not s:
                return ""
            # Replace common unicode chars to match latin-1 encoding safely
            return s.encode('latin-1', 'replace').decode('latin-1')

        # Title
        pdf.set_font("Helvetica", size=20, style="B")
        pdf.cell(0, 10, text=safe_txt(page["title"]), new_x="LMARGIN", new_y="NEXT", align="L")
        pdf.ln(5)

        # Summary
        if page.get("summary"):
            pdf.set_font("Helvetica", size=11, style="I")
            pdf.multi_cell(0, 7, text=safe_txt(page["summary"]))
            pdf.ln(5)

        # Content
        if page["content"]:
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 6, text=safe_txt(page["content"]))
            pdf.ln(10)

        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
        pdf.ln(5)

        # Section Header
        pdf.set_font("Helvetica", size=14, style="B")
        pdf.cell(0, 10, text="Referenced Notes & Insights", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

        for idx, item in enumerate(items, 1):
            pdf.set_font("Helvetica", size=12, style="B")
            pdf.cell(0, 8, text=safe_txt(f"{idx}. {item['title']}"), new_x="LMARGIN", new_y="NEXT")
            
            pdf.set_font("Helvetica", size=9, style="I")
            pdf.cell(0, 5, text=safe_txt(f"Type: {item['type'].capitalize()}"), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 6, text=safe_txt(item["content"]))
            pdf.ln(5)
            
            # Divider line
            if idx < len(items):
                pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
                pdf.ln(5)

        # Output bytes
        pdf_bytes = pdf.output()
        return pdf_bytes if isinstance(pdf_bytes, bytes) else bytes(pdf_bytes)

    @staticmethod
    def export_docx(page_id: str, context: WorkspaceContext) -> bytes:
        """Constructs a DOCX document using python-docx."""
        page = KnowledgePageService._verify_page_ownership(page_id, context)
        items = KnowledgePageService.get_references(page_id, context)

        doc = DocxDocument()
        
        # Title
        doc.add_heading(page["title"], level=0)
        
        # Summary
        if page.get("summary"):
            doc.add_paragraph(page["summary"], style="Subtitle")
            
        # Content
        if page["content"]:
            doc.add_paragraph(page["content"])
            
        # References header
        doc.add_heading("Referenced Notes & Insights", level=1)
        
        for idx, item in enumerate(items, 1):
            doc.add_heading(f"{idx}. {item['title']}", level=2)
            doc.add_paragraph(f"Type: {item['type'].capitalize()}", style="Normal").runs[0].italic = True
            doc.add_paragraph(item["content"])
            
            if idx < len(items):
                doc.add_paragraph("─" * 40)

        # Save to buffer
        stream = BytesIO()
        doc.save(stream)
        return stream.getvalue()
